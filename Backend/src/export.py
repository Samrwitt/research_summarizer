# src/export.py
from __future__ import annotations

from typing import Dict, Any, List
from datetime import datetime
from docx import Document
from docx.shared import Pt


def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_codeblock(doc: Document, text: str):
    # DOCX doesn't have native code blocks; use monospace-ish styling.
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(6)
    return p


def save_preprocess_docx(pre: Dict[str, Any], out_path: str) -> str:
    """
    Save preprocess output into a human-readable DOCX.
    """
    doc = Document()

    title = pre.get("title") or "(Untitled)"
    paper_id = pre.get("paper_id")
    source = pre.get("source")
    meta = pre.get("meta") or {}
    stats = pre.get("stats") or {}

    _add_heading(doc, "Preprocess Output", level=0)
    _add_paragraph(doc, f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    _add_paragraph(doc, f"Title: {title}")
    _add_paragraph(doc, f"Source: {source}" + (f" | Paper ID: {paper_id}" if paper_id else ""))

    if meta.get("abs_url"):
        _add_paragraph(doc, f"arXiv URL: {meta.get('abs_url')}")
    if meta.get("html_url"):
        _add_paragraph(doc, f"HTML URL: {meta.get('html_url')}")
    if meta.get("pdf_url"):
        _add_paragraph(doc, f"PDF URL: {meta.get('pdf_url')}")

    doc.add_paragraph("")

    # Abstract
    abstract = pre.get("abstract")
    _add_heading(doc, "Abstract", level=1)
    _add_paragraph(doc, abstract.strip() if abstract else "(No abstract found)")

    # Stats
    _add_heading(doc, "Stats", level=1)
    for k, v in stats.items():
        _add_paragraph(doc, f"{k}: {v}")

    # Sections
    sections: Dict[str, str] = pre.get("sections") or {}
    _add_heading(doc, "Detected Sections", level=1)
    if not sections:
        _add_paragraph(doc, "(No sections detected)")
    else:
        for name, content in sections.items():
            _add_heading(doc, name.title(), level=2)
            _add_paragraph(doc, content[:2000] + ("…" if len(content) > 2000 else ""))

    # Focus text (preview)
    _add_heading(doc, "Focus Text (Preview)", level=1)
    focus_text = pre.get("focus_text") or ""
    _add_paragraph(doc, focus_text[:3000] + ("…" if len(focus_text) > 3000 else ""))

    # Chunks (preview)
    chunks: List[str] = pre.get("chunks") or []
    _add_heading(doc, f"Chunks (count={len(chunks)})", level=1)
    for i, ch in enumerate(chunks[:10], start=1):  # only first 10 by default
        _add_heading(doc, f"Chunk {i}", level=2)
        _add_codeblock(doc, ch[:2500] + ("…" if len(ch) > 2500 else ""))

    doc.save(out_path)
    return out_path
