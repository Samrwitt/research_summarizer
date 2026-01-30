import os
from docx import Document

def export_markdown(md_content, output_path):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

def export_docx(data, summary_text, bullets, output_path):
    doc = Document()
    
    title = data.get('title', 'Paper Summary')
    doc.add_heading(title, 0)
    
    doc.add_paragraph(f"Source: {data['source']}")
    if data.get('paper_id'):
        doc.add_paragraph(f"ID: {data['paper_id']}")
    
    doc.add_heading('TL;DR', level=1)
    doc.add_paragraph(summary_text)
    
    doc.add_heading('Key Points', level=1)
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
        
    doc.save(output_path)
